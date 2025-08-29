#!/usr/bin/env python3
"""
Word Document Generator using python-docx and docxtpl
Demonstrates contract generation, report creation, and dynamic content insertion
"""

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime, timedelta
import json
import os
from typing import Dict, List, Any, Optional

class WordDocumentGenerator:
    """Main class for generating Word documents with templates"""
    
    def __init__(self, templates_dir: str = "./templates"):
        self.templates_dir = templates_dir
        self.ensure_templates_dir()
    
    def ensure_templates_dir(self):
        """Create templates directory if it doesn't exist"""
        if not os.path.exists(self.templates_dir):
            os.makedirs(self.templates_dir)
            print(f"Created templates directory: {self.templates_dir}")
    
    def generate_contract(self, contract_data: Dict[str, Any], output_path: str = None) -> str:
        """
        Generate a contract document using template-based approach
        
        Args:
            contract_data: Dictionary containing contract information
            output_path: Output file path (optional)
        
        Returns:
            Path to generated document
        """
        # Create template if it doesn't exist
        template_path = os.path.join(self.templates_dir, "contract_template.docx")
        if not os.path.exists(template_path):
            self.create_contract_template(template_path)
        
        # Load template
        doc = DocxTemplate(template_path)
        
        # Prepare context data
        context = self.prepare_contract_context(contract_data)
        
        # Render template
        doc.render(context)
        
        # Save document
        if output_path is None:
            client_name = contract_data.get('client_name', 'client').replace(' ', '_')
            output_path = f"contract_{client_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc.save(output_path)
        print(f"Contract generated: {output_path}")
        return output_path
    
    def create_contract_template(self, template_path: str):
        """Create a contract template document"""
        doc = Document()
        
        # Title
        title = doc.add_heading('SERVICE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contract details section
        doc.add_heading('Contract Details', level=1)
        
        # Create table for contract details
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'
        
        # Add template placeholders
        details = [
            ('Client Name', '{{client_name}}'),
            ('Project Name', '{{project_name}}'),
            ('Start Date', '{{start_date}}'),
            ('End Date', '{{end_date}}'),
            ('Total Amount', '{{total_amount}}'),
            ('Contract Date', '{{contract_date}}')
        ]
        
        for field, value in details:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = value
        
        # Terms and conditions
        doc.add_heading('Terms and Conditions', level=1)
        doc.add_paragraph('{{terms}}')
        
        # Milestones
        doc.add_heading('Project Milestones', level=1)
        doc.add_paragraph('{{milestones}}')
        
        # Signature section
        doc.add_heading('Signatures', level=1)
        doc.add_paragraph('Client Signature: _________________')
        doc.add_paragraph('Date: {{client_signature_date}}')
        doc.add_paragraph('')
        doc.add_paragraph('Service Provider Signature: _________________')
        doc.add_paragraph('Date: {{provider_signature_date}}')
        
        # Save template
        doc.save(template_path)
        print(f"Contract template created: {template_path}")
    
    def prepare_contract_context(self, contract_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context data for contract template"""
        start_date = datetime.strptime(contract_data['start_date'], '%Y-%m-%d')
        end_date = datetime.strptime(contract_data['end_date'], '%Y-%m-%d')
        
        context = {
            'client_name': contract_data['client_name'],
            'project_name': contract_data['project_name'],
            'start_date': start_date.strftime('%B %d, %Y'),
            'end_date': end_date.strftime('%B %d, %Y'),
            'total_amount': f"${contract_data['amount']:,.2f}",
            'contract_date': datetime.now().strftime('%B %d, %Y'),
            'terms': contract_data.get('terms', 'Standard terms and conditions apply.'),
            'milestones': self.format_milestones(contract_data.get('milestones', [])),
            'client_signature_date': datetime.now().strftime('%B %d, %Y'),
            'provider_signature_date': datetime.now().strftime('%B %d, %Y')
        }
        
        return context
    
    def format_milestones(self, milestones: List[Dict[str, Any]]) -> str:
        """Format milestones for template insertion"""
        if not milestones:
            return "No specific milestones defined."
        
        formatted = []
        for i, milestone in enumerate(milestones, 1):
            formatted.append(f"{i}. {milestone['description']} - Due: {milestone['due_date']}")
        
        return '\n'.join(formatted)
    
    def generate_report(self, report_data: Dict[str, Any], output_path: str = None) -> str:
        """
        Generate a report document programmatically
        
        Args:
            report_data: Dictionary containing report information
            output_path: Output file path (optional)
        
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading(report_data.get('title', 'Project Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(report_data.get('summary', 'No summary provided.'))
        
        # Project Details
        doc.add_heading('Project Details', level=1)
        self.add_project_details_table(doc, report_data.get('project_details', {}))
        
        # Metrics
        if 'metrics' in report_data:
            doc.add_heading('Key Metrics', level=1)
            self.add_metrics_table(doc, report_data['metrics'])
        
        # Financial Data
        if 'financials' in report_data and report_data.get('include_financials', False):
            doc.add_heading('Financial Summary', level=1)
            self.add_financial_table(doc, report_data['financials'])
        
        # Risk Assessment
        if report_data.get('status') == 'At Risk':
            doc.add_heading('Risk Assessment', level=1)
            self.add_risk_section(doc, report_data.get('risk_plan', []))
        
        # Recommendations
        if 'recommendations' in report_data:
            doc.add_heading('Recommendations', level=1)
            self.add_recommendations_list(doc, report_data['recommendations'])
        
        # Save document
        if output_path is None:
            project_name = report_data.get('project_name', 'project').replace(' ', '_')
            output_path = f"report_{project_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc.save(output_path)
        print(f"Report generated: {output_path}")
        return output_path
    
    def add_project_details_table(self, doc: Document, details: Dict[str, Any]):
        """Add project details table to document"""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Add data rows
        for metric, value in details.items():
            row_cells = table.add_row().cells
            row_cells[0].text = metric.replace('_', ' ').title()
            row_cells[1].text = str(value)
    
    def add_metrics_table(self, doc: Document, metrics: Dict[str, Any]):
        """Add metrics table to document"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Current Value'
        hdr_cells[2].text = 'Target'
        
        # Add data rows
        for metric, data in metrics.items():
            row_cells = table.add_row().cells
            row_cells[0].text = metric.replace('_', ' ').title()
            row_cells[1].text = str(data.get('current', 'N/A'))
            row_cells[2].text = str(data.get('target', 'N/A'))
    
    def add_financial_table(self, doc: Document, financials: List[Dict[str, Any]]):
        """Add financial data table to document"""
        if not financials:
            doc.add_paragraph('No financial data available.')
            return
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Period'
        hdr_cells[1].text = 'Budget'
        hdr_cells[2].text = 'Actual'
        hdr_cells[3].text = 'Variance'
        
        # Add data rows
        for item in financials:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('period', 'N/A')
            row_cells[1].text = f"${item.get('budget', 0):,.2f}"
            row_cells[2].text = f"${item.get('actual', 0):,.2f}"
            
            variance = item.get('actual', 0) - item.get('budget', 0)
            variance_text = f"${variance:,.2f}"
            if variance < 0:
                variance_text = f"({variance_text})"
            row_cells[3].text = variance_text
    
    def add_risk_section(self, doc: Document, risk_plan: List[Dict[str, Any]]):
        """Add risk assessment section to document"""
        doc.add_paragraph('The following risks have been identified and mitigation strategies are in place:')
        
        for risk in risk_plan:
            p = doc.add_paragraph()
            p.add_run(f"Risk: {risk.get('description', 'N/A')}").bold = True
            doc.add_paragraph(f"Impact: {risk.get('impact', 'N/A')}")
            doc.add_paragraph(f"Mitigation: {risk.get('mitigation', 'N/A')}")
            doc.add_paragraph('')  # Add spacing
    
    def add_recommendations_list(self, doc: Document, recommendations: List[str]):
        """Add recommendations list to document"""
        for i, recommendation in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {recommendation}", style='List Number')
    
    def generate_letter(self, letter_data: Dict[str, Any], output_path: str = None) -> str:
        """
        Generate a business letter
        
        Args:
            letter_data: Dictionary containing letter information
            output_path: Output file path (optional)
        
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Letter header
        header = doc.add_paragraph()
        header.add_run(letter_data.get('company_name', 'Company Name')).bold = True
        header.add_run(f"\n{letter_data.get('company_address', 'Company Address')}")
        
        # Date
        doc.add_paragraph(f"\n{datetime.now().strftime('%B %d, %Y')}")
        
        # Recipient
        doc.add_paragraph(f"\n{letter_data.get('recipient_name', 'Recipient Name')}")
        doc.add_paragraph(letter_data.get('recipient_address', 'Recipient Address'))
        
        # Salutation
        doc.add_paragraph(f"\nDear {letter_data.get('recipient_name', 'Recipient')},")
        
        # Letter body
        doc.add_paragraph(letter_data.get('body', 'Letter body content.'))
        
        # Closing
        doc.add_paragraph(f"\nSincerely,\n{letter_data.get('sender_name', 'Sender Name')}")
        
        # Save document
        if output_path is None:
            recipient_name = letter_data.get('recipient_name', 'recipient').replace(' ', '_')
            output_path = f"letter_{recipient_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc.save(output_path)
        print(f"Letter generated: {output_path}")
        return output_path

def create_sample_documents():
    """Create sample documents to demonstrate the generator"""
    generator = WordDocumentGenerator()
    
    # Sample contract data
    contract_data = {
        'client_name': 'Acme Corporation',
        'project_name': 'Website Redesign Project',
        'start_date': '2024-01-15',
        'end_date': '2024-06-15',
        'amount': 50000.00,
        'terms': 'This agreement outlines the terms and conditions for the website redesign project. The service provider will deliver a modern, responsive website that meets all specified requirements.',
        'milestones': [
            {
                'description': 'Design approval and wireframes',
                'due_date': '2024-02-15'
            },
            {
                'description': 'Frontend development completion',
                'due_date': '2024-04-15'
            },
            {
                'description': 'Backend integration and testing',
                'due_date': '2024-05-15'
            },
            {
                'description': 'Final delivery and launch',
                'due_date': '2024-06-15'
            }
        ]
    }
    
    # Sample report data
    report_data = {
        'title': 'Q1 2024 Project Status Report',
        'project_name': 'E-commerce Platform Development',
        'summary': 'The e-commerce platform development project is progressing well with 75% completion. All major milestones have been met on schedule, and the team is on track for the planned launch date.',
        'project_details': {
            'Project Manager': 'Sarah Johnson',
            'Team Size': '8 developers',
            'Start Date': 'January 1, 2024',
            'Planned End Date': 'June 30, 2024',
            'Current Status': 'On Track',
            'Completion Percentage': '75%'
        },
        'metrics': {
            'code_coverage': {'current': 85, 'target': 90},
            'test_pass_rate': {'current': 98, 'target': 95},
            'bug_count': {'current': 12, 'target': 20},
            'user_stories_completed': {'current': 45, 'target': 50}
        },
        'include_financials': True,
        'financials': [
            {'period': 'January', 'budget': 50000, 'actual': 48000},
            {'period': 'February', 'budget': 50000, 'actual': 52000},
            {'period': 'March', 'budget': 50000, 'actual': 49000}
        ],
        'status': 'On Track',
        'recommendations': [
            'Increase code coverage to meet target of 90%',
            'Implement additional automated testing',
            'Schedule user acceptance testing for next month',
            'Prepare deployment documentation'
        ]
    }
    
    # Sample letter data
    letter_data = {
        'company_name': 'TechCorp Solutions',
        'company_address': '123 Business Ave\nSan Francisco, CA 94105',
        'recipient_name': 'John Smith',
        'recipient_address': '456 Corporate Blvd\nNew York, NY 10001',
        'sender_name': 'Jane Doe',
        'body': 'I am writing to follow up on our recent discussion regarding the potential partnership between our companies. We believe there are significant opportunities for collaboration that would benefit both organizations.\n\nWe would like to schedule a meeting to discuss the details further and explore how we can work together to achieve our mutual goals.'
    }
    
    try:
        # Generate documents
        print("Generating sample documents...")
        
        contract_path = generator.generate_contract(contract_data)
        report_path = generator.generate_report(report_data)
        letter_path = generator.generate_letter(letter_data)
        
        print(f"\nSample documents generated successfully:")
        print(f"- Contract: {contract_path}")
        print(f"- Report: {report_path}")
        print(f"- Letter: {letter_path}")
        
    except Exception as e:
        print(f"Error generating documents: {e}")

if __name__ == "__main__":
    create_sample_documents()
